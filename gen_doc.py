# -*- coding: utf-8 -*-
import json
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# 读取内容
with open(r'E:\项目\lmg-platform\doc_content.json', 'r', encoding='utf-8') as f:
    data = json.load(f)

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5

# 标题
title = doc.add_heading(data['title'], level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run(data['subtitle'])
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_page_break()

# 目录
doc.add_heading('目录', level=1)
for item in data['toc']:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

sections = data['sections']

# ===== 1. 项目概述 =====
s = sections['1']
doc.add_heading(s['title'], level=1)
doc.add_paragraph(s['content'])
doc.add_paragraph('核心功能：')
for f in s['features']:
    doc.add_paragraph(f, style='List Bullet')

# ===== 2. 技术栈 =====
s = sections['2']
doc.add_heading(s['title'], level=1)
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = '层级'
hdr[1].text = '技术'
hdr[2].text = '说明'
for layer, tech, desc in s['techs']:
    row = table.add_row().cells
    row[0].text = layer
    row[1].text = tech
    row[2].text = desc

# ===== 3. 项目结构 =====
s = sections['3']
doc.add_heading(s['title'], level=1)

doc.add_heading('3.1 后端目录结构', level=2)
for path, desc in s['backend_files']:
    p = doc.add_paragraph()
    run = p.add_run(path)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.bold = True
    p.add_run('\n    ' + desc)

doc.add_heading('3.2 前端目录结构', level=2)
for path, desc in s['frontend_files']:
    p = doc.add_paragraph()
    run = p.add_run(path)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.bold = True
    p.add_run('\n    ' + desc)

# ===== 4. 数据源 =====
s = sections['4']
doc.add_heading(s['title'], level=1)
doc.add_paragraph(s['intro'])

ds_table = doc.add_table(rows=1, cols=4)
ds_table.style = 'Light Grid Accent 1'
ds_hdr = ds_table.rows[0].cells
ds_hdr[0].text = '文件'
ds_hdr[1].text = '类型'
ds_hdr[2].text = '说明'
ds_hdr[3].text = '时间字段'
for file, typ, desc, tf in s['sources']:
    row = ds_table.add_row().cells
    row[0].text = file
    row[1].text = typ
    row[2].text = desc
    row[3].text = tf

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('重要说明：')
run.font.bold = True
doc.add_paragraph(s['note'], style='List Bullet')

# ===== 5. 数据库 =====
s = sections['5']
doc.add_heading(s['title'], level=1)
doc.add_paragraph(s['intro'])

for name, desc, fields in s['tables']:
    p = doc.add_paragraph()
    run = p.add_run(name)
    run.font.bold = True
    run.font.name = 'Consolas'
    p.add_run(' - ' + desc)
    p2 = doc.add_paragraph('    字段: ' + fields)
    p2.paragraph_format.space_after = Pt(2)

# ===== 6. 利润算法 =====
s = sections['6']
doc.add_heading(s['title'], level=1)

doc.add_heading('6.1 净利润公式', level=2)
doc.add_paragraph(s['formula'])
doc.add_paragraph('其中：')
for item in s['formula_items']:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('6.2 净销量算法', level=2)
doc.add_paragraph(s['net_sales'])

doc.add_heading('6.3 无订单产品的处理', level=2)
doc.add_paragraph(s['no_order'])

doc.add_heading('6.4 数据查询逻辑', level=2)
doc.add_paragraph(s['query_logic'])

# ===== 7. API =====
s = sections['7']
doc.add_heading(s['title'], level=1)

api_table = doc.add_table(rows=1, cols=4)
api_table.style = 'Light Grid Accent 1'
api_hdr = api_table.rows[0].cells
api_hdr[0].text = '方法'
api_hdr[1].text = '路径'
api_hdr[2].text = '参数'
api_hdr[3].text = '说明'
for method, path, params, desc in s['apis']:
    row = api_table.add_row().cells
    row[0].text = method
    row[1].text = path
    row[2].text = params
    row[3].text = desc

# ===== 8. 问题 =====
s = sections['8']
doc.add_heading(s['title'], level=1)

for prob in s['problems']:
    doc.add_heading(prob['title'], level=2)

    p = doc.add_paragraph()
    run = p.add_run('现象：')
    run.font.bold = True
    p.add_run(prob['symptom'])

    p = doc.add_paragraph()
    run = p.add_run('原因：')
    run.font.bold = True
    p.add_run(prob['cause'])

    p = doc.add_paragraph()
    run = p.add_run('文件：')
    run.font.bold = True
    run2 = p.add_run(prob['file'])
    run2.font.name = 'Consolas'
    run2.font.size = Pt(10)

    p = doc.add_paragraph()
    run = p.add_run('解决方案：')
    run.font.bold = True
    p.add_run(prob['fix'])

# ===== 9. 部署 =====
s = sections['9']
doc.add_heading(s['title'], level=1)

doc.add_heading('9.1 启动命令', level=2)
for name, cmd in s['commands']:
    p = doc.add_paragraph()
    run = p.add_run(name + ': ')
    run.font.bold = True
    run2 = p.add_run(cmd)
    run2.font.name = 'Consolas'
    run2.font.size = Pt(10)

doc.add_heading('9.2 访问地址', level=2)
for addr in s['addresses']:
    doc.add_paragraph(addr, style='List Bullet')

doc.add_heading('9.3 防火墙配置', level=2)
doc.add_paragraph('以管理员身份运行 CMD：')
p = doc.add_paragraph()
run = p.add_run(s['firewall'])
run.font.name = 'Consolas'
run.font.size = Pt(10)

# 保存
output_path = os.path.expanduser('~/Desktop/LMG数据平台技术文档.docx')
doc.save(output_path)
print('Done: ' + output_path)
